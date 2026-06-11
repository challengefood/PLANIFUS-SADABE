import base64
import calendar
import io
import os
import re
import sqlite3
import unicodedata
from datetime import date, datetime
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

import altair as alt
import pandas as pd
import streamlit as st

try:
    from docx import Document
except Exception:
    Document = None

APP_NAME = "SADABE PLANIFIUS"
BASE_DIR = Path(__file__).resolve().parent
ASSETS_DIR = BASE_DIR / "assets"
SAMPLE_DIR = BASE_DIR / "sample_data"
LOGO_PATH = ASSETS_DIR / "logo_sadabe.png"
DB_PATH = Path(os.getenv("SADABE_PLANIFIUS_DB", BASE_DIR / "sadabe_planifius.db"))

DEFAULT_PROJECTS = [
    ("SOS Lemurs", "Projet de conservation et suivi des lémuriens"),
    ("Darwin Initiatives", "Projet / financement Darwin Initiatives"),
    ("Seacology", "Projet / financement Seacology"),
    ("Rainforest Trust", "Projet / financement Rainforest Trust"),
]
DEFAULT_PARTNERS = [
    ("TGBS (MBG)", ""),
    ("MfM", ""),
    ("UWE", ""),
    ("Regen", ""),
    ("UNI", ""),
    ("ENS", ""),
]
PRIORITIES = ["Faible", "Moyenne", "Haute", "Critique"]
STATUSES = ["À faire", "En cours", "Terminé", "Reporté", "Annulé"]
DONE_STATUSES = {"Terminé", "Annulé"}


# -----------------------------------------------------------------------------
# UI helpers
# -----------------------------------------------------------------------------

def image_to_base64(path: Path) -> str:
    if not path.exists():
        return ""
    return base64.b64encode(path.read_bytes()).decode("utf-8")


def inject_css() -> None:
    st.markdown(
        """
        <style>
        :root {
            --sadabe-green:#2E7D32;
            --sadabe-dark:#1B5E20;
            --sadabe-soft:#E8F5E9;
            --sadabe-gold:#B7791F;
        }
        .main .block-container { padding-top: 1.2rem; max-width: 1320px; }
        .sadabe-header {
            display:flex; gap:1.2rem; align-items:center;
            padding: 1rem 1.2rem; border-radius: 18px;
            background: linear-gradient(135deg, #E8F5E9 0%, #FFFFFF 70%);
            border: 1px solid #D9EAD3; margin-bottom: 1rem;
        }
        .sadabe-title { font-size: 2.1rem; font-weight: 800; color: var(--sadabe-dark); margin:0; }
        .sadabe-subtitle { font-size: .98rem; color:#38523A; margin-top:.25rem; }
        .metric-card {
            border-radius: 16px; padding: 1rem; background: #FFFFFF;
            border: 1px solid #E1E7E1; box-shadow: 0 2px 10px rgba(27,94,32,.06);
        }
        .metric-label { color:#5f6b61; font-size:.85rem; }
        .metric-value { color:#1B5E20; font-size:1.85rem; font-weight:800; }
        .danger { color:#B91C1C; font-weight:700; }
        .warning { color:#B7791F; font-weight:700; }
        .ok { color:#2E7D32; font-weight:700; }
        .small-note { color:#667; font-size:.85rem; }
        div[data-testid="stSidebar"] { background: #F6FAF6; }
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_header() -> None:
    logo_b64 = image_to_base64(LOGO_PATH)
    logo_html = f'<img src="data:image/png;base64,{logo_b64}" style="width:90px; height:auto;">' if logo_b64 else ""
    st.markdown(
        f"""
        <div class="sadabe-header">
            <div>{logo_html}</div>
            <div>
                <p class="sadabe-title">{APP_NAME}</p>
                <div class="sadabe-subtitle">
                    Dashboard de planification mensuelle : activités, responsables, équipe, projets et partenaires.
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def kpi_card(label: str, value: str, note: str = "") -> None:
    st.markdown(
        f"""
        <div class="metric-card">
            <div class="metric-label">{label}</div>
            <div class="metric-value">{value}</div>
            <div class="small-note">{note}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def month_label(month_value: str) -> str:
    try:
        year, month = map(int, month_value.split("-"))
        return f"{calendar.month_name[month].capitalize()} {year}"
    except Exception:
        return month_value


# -----------------------------------------------------------------------------
# Database
# -----------------------------------------------------------------------------

@st.cache_resource
def get_connection() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    init_db(conn)
    return conn


def init_db(conn: sqlite3.Connection) -> None:
    cur = conn.cursor()
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS projects (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            description TEXT DEFAULT '',
            active INTEGER DEFAULT 1,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS partners (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            description TEXT DEFAULT '',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS team_members (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            poste TEXT DEFAULT '',
            email TEXT DEFAULT '',
            phone TEXT DEFAULT '',
            active INTEGER DEFAULT 1,
            notes TEXT DEFAULT '',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS activities (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            description TEXT DEFAULT '',
            project_id INTEGER,
            partner_id INTEGER,
            period_month TEXT DEFAULT '',
            planned_date TEXT DEFAULT '',
            deadline TEXT DEFAULT '',
            priority TEXT DEFAULT 'Moyenne',
            status TEXT DEFAULT 'À faire',
            responsible_id INTEGER,
            location TEXT DEFAULT '',
            expected_output TEXT DEFAULT '',
            notes TEXT DEFAULT '',
            source_file TEXT DEFAULT '',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(project_id) REFERENCES projects(id),
            FOREIGN KEY(partner_id) REFERENCES partners(id),
            FOREIGN KEY(responsible_id) REFERENCES team_members(id)
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS activity_members (
            activity_id INTEGER NOT NULL,
            member_id INTEGER NOT NULL,
            role_in_activity TEXT DEFAULT '',
            PRIMARY KEY(activity_id, member_id),
            FOREIGN KEY(activity_id) REFERENCES activities(id) ON DELETE CASCADE,
            FOREIGN KEY(member_id) REFERENCES team_members(id) ON DELETE CASCADE
        )
        """
    )
    for name, desc in DEFAULT_PROJECTS:
        cur.execute("INSERT OR IGNORE INTO projects(name, description, active) VALUES (?, ?, 1)", (name, desc))
    for name, desc in DEFAULT_PARTNERS:
        cur.execute("INSERT OR IGNORE INTO partners(name, description) VALUES (?, ?)", (name, desc))
    conn.commit()


def qdf(sql: str, params: Tuple = ()) -> pd.DataFrame:
    return pd.read_sql_query(sql, get_connection(), params=params)


def execute(sql: str, params: Tuple = ()) -> sqlite3.Cursor:
    conn = get_connection()
    cur = conn.cursor()
    cur.execute(sql, params)
    conn.commit()
    return cur


def fetch_one(sql: str, params: Tuple = ()) -> Optional[sqlite3.Row]:
    return get_connection().execute(sql, params).fetchone()


def list_projects(active_only: bool = False) -> pd.DataFrame:
    where = "WHERE active = 1" if active_only else ""
    return qdf(f"SELECT id, name, description, active FROM projects {where} ORDER BY name")


def list_partners() -> pd.DataFrame:
    return qdf("SELECT id, name, description FROM partners ORDER BY name")


def list_team(active_only: bool = False) -> pd.DataFrame:
    where = "WHERE active = 1" if active_only else ""
    return qdf(f"SELECT id, name, poste, email, phone, active, notes FROM team_members {where} ORDER BY name")


def ensure_project(name: str, description: str = "") -> Optional[int]:
    name = clean_text(name)
    if not name:
        return None
    execute("INSERT OR IGNORE INTO projects(name, description, active) VALUES (?, ?, 1)", (name, description))
    row = fetch_one("SELECT id FROM projects WHERE name = ?", (name,))
    return int(row["id"]) if row else None


def ensure_partner(name: str, description: str = "") -> Optional[int]:
    name = clean_text(name)
    if not name:
        return None
    execute("INSERT OR IGNORE INTO partners(name, description) VALUES (?, ?)", (name, description))
    row = fetch_one("SELECT id FROM partners WHERE name = ?", (name,))
    return int(row["id"]) if row else None


def ensure_member(name: str, poste: str = "") -> Optional[int]:
    name = clean_text(name)
    if not name:
        return None
    execute(
        "INSERT OR IGNORE INTO team_members(name, poste, active) VALUES (?, ?, 1)",
        (name, clean_text(poste)),
    )
    row = fetch_one("SELECT id FROM team_members WHERE name = ?", (name,))
    return int(row["id"]) if row else None


# -----------------------------------------------------------------------------
# Normalization / import helpers
# -----------------------------------------------------------------------------

def clean_text(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and pd.isna(value):
        return ""
    return str(value).strip()


def normalize(value: str) -> str:
    value = clean_text(value).lower()
    value = unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode("ascii")
    value = re.sub(r"[^a-z0-9]+", " ", value)
    return re.sub(r"\s+", " ", value).strip()


COLUMN_SYNONYMS = {
    "title": ["activite", "activité", "activity", "tache", "tâche", "task", "titre", "action", "libelle", "libellé"],
    "description": ["description", "details", "detail", "détails", "commentaire"],
    "project": ["projet", "project", "programme", "bailleur", "financement"],
    "partner": ["partenaire", "partner", "partenaires", "collaborateur"],
    "period_month": ["mois prevu", "mois prévu", "mois", "periode", "période", "period", "month", "yyyymm", "yyyy mm"],
    "planned_date": ["jour prevu", "jour prévu", "date prevue", "date prévue", "date planifiee", "date planifiée", "planned date", "date"],
    "deadline": ["date limite", "deadline", "echeance", "échéance", "due date", "limite"],
    "priority": ["priorite", "priorité", "priority", "urgence", "urgent"],
    "status": ["statut", "status", "etat", "état", "avancement"],
    "responsible": ["responsable", "responsable principal", "owner", "lead", "charge", "chargé"],
    "members": ["membres", "membres equipe", "membres équipe", "equipe", "équipe", "participants", "staff"],
    "location": ["lieu", "site", "localisation", "village", "zone"],
    "expected_output": ["resultat attendu", "résultat attendu", "output", "livrable", "produit", "deliverable"],
    "notes": ["notes", "observation", "observations", "remarque", "remarks"],
}


def build_column_map(columns: Iterable[str]) -> Dict[str, str]:
    normalized_columns = {normalize(c): c for c in columns}
    mapping = {}
    for key, synonyms in COLUMN_SYNONYMS.items():
        for synonym in synonyms:
            norm_syn = normalize(synonym)
            if norm_syn in normalized_columns:
                mapping[key] = normalized_columns[norm_syn]
                break
        if key not in mapping:
            for norm_col, original in normalized_columns.items():
                if any(normalize(synonym) in norm_col for synonym in synonyms):
                    mapping[key] = original
                    break
    return mapping


def get_col_value(row: pd.Series, mapping: Dict[str, str], key: str, default=""):
    col = mapping.get(key)
    if not col:
        return default
    value = row.get(col, default)
    if isinstance(value, float) and pd.isna(value):
        return default
    return value


def parse_date_value(value) -> str:
    value = clean_text(value)
    if not value:
        return ""
    try:
        parsed = pd.to_datetime(value, dayfirst=True, errors="coerce")
        if pd.isna(parsed):
            return ""
        return parsed.date().isoformat()
    except Exception:
        return ""


def parse_month_value(value, fallback_date: str = "", default_month: str = "") -> str:
    raw = clean_text(value)
    if raw:
        if re.match(r"^\d{4}-\d{2}$", raw):
            return raw
        parsed_date = parse_date_value(raw)
        if parsed_date:
            return parsed_date[:7]
        month_match = re.search(r"(20\d{2})[-_/ ]?(0?[1-9]|1[0-2])", raw)
        if month_match:
            return f"{month_match.group(1)}-{int(month_match.group(2)):02d}"
    if fallback_date:
        return fallback_date[:7]
    return default_month


def normalize_priority(value: str) -> str:
    text = normalize(value)
    if any(k in text for k in ["crit", "tres urgent", "très urgent", "immediate", "immediat"]):
        return "Critique"
    if any(k in text for k in ["haut", "haute", "high", "urgent", "fort"]):
        return "Haute"
    if any(k in text for k in ["faible", "low", "bas"]):
        return "Faible"
    return "Moyenne"


def normalize_status(value: str) -> str:
    text = normalize(value)
    if any(k in text for k in ["termine", "done", "completed", "realise", "réalisé"]):
        return "Terminé"
    if any(k in text for k in ["cours", "progress", "ongoing"]):
        return "En cours"
    if any(k in text for k in ["reporte", "report", "postpone"]):
        return "Reporté"
    if any(k in text for k in ["annule", "cancel"]):
        return "Annulé"
    return "À faire"


def split_members(value: str) -> List[str]:
    text = clean_text(value)
    if not text:
        return []
    return [m.strip() for m in re.split(r"[;,\n]+", text) if m.strip()]


def urgency_for_row(row: pd.Series) -> str:
    status = clean_text(row.get("Statut", ""))
    priority = clean_text(row.get("Priorité", ""))
    deadline = parse_date_value(row.get("Date limite", ""))
    if status in DONE_STATUSES:
        return "Clôturé"
    if not deadline:
        return "À planifier" if not clean_text(row.get("Jour prévu", "")) else "Normal"
    today = date.today()
    try:
        d = datetime.fromisoformat(deadline).date()
    except Exception:
        return "Normal"
    days = (d - today).days
    if days < 0:
        return "En retard"
    if priority in {"Critique", "Haute"} or days <= 7:
        return "Urgent"
    if days <= 14:
        return "Bientôt"
    return "Normal"


def days_to_deadline(row: pd.Series) -> Optional[int]:
    deadline = parse_date_value(row.get("Date limite", ""))
    if not deadline:
        return None
    try:
        return (datetime.fromisoformat(deadline).date() - date.today()).days
    except Exception:
        return None


# -----------------------------------------------------------------------------
# Activities CRUD / queries
# -----------------------------------------------------------------------------

def add_activity(
    title: str,
    description: str,
    project_id: Optional[int],
    partner_id: Optional[int],
    period_month: str,
    planned_date: str,
    deadline: str,
    priority: str,
    status: str,
    responsible_id: Optional[int],
    location: str,
    expected_output: str,
    notes: str,
    member_roles: Dict[int, str],
    source_file: str = "",
) -> int:
    cur = execute(
        """
        INSERT INTO activities(
            title, description, project_id, partner_id, period_month, planned_date, deadline,
            priority, status, responsible_id, location, expected_output, notes, source_file, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
        """,
        (
            title,
            description,
            project_id,
            partner_id,
            period_month,
            planned_date,
            deadline,
            priority,
            status,
            responsible_id,
            location,
            expected_output,
            notes,
            source_file,
        ),
    )
    activity_id = int(cur.lastrowid)
    for member_id, role in member_roles.items():
        execute(
            "INSERT OR IGNORE INTO activity_members(activity_id, member_id, role_in_activity) VALUES (?, ?, ?)",
            (activity_id, int(member_id), clean_text(role)),
        )
    return activity_id


def update_activity(activity_id: int, data: dict, member_roles: Dict[int, str]) -> None:
    execute(
        """
        UPDATE activities SET
            title=?, description=?, project_id=?, partner_id=?, period_month=?, planned_date=?, deadline=?,
            priority=?, status=?, responsible_id=?, location=?, expected_output=?, notes=?, updated_at=CURRENT_TIMESTAMP
        WHERE id=?
        """,
        (
            data["title"],
            data.get("description", ""),
            data.get("project_id"),
            data.get("partner_id"),
            data.get("period_month", ""),
            data.get("planned_date", ""),
            data.get("deadline", ""),
            data.get("priority", "Moyenne"),
            data.get("status", "À faire"),
            data.get("responsible_id"),
            data.get("location", ""),
            data.get("expected_output", ""),
            data.get("notes", ""),
            int(activity_id),
        ),
    )
    execute("DELETE FROM activity_members WHERE activity_id=?", (int(activity_id),))
    for member_id, role in member_roles.items():
        execute(
            "INSERT OR IGNORE INTO activity_members(activity_id, member_id, role_in_activity) VALUES (?, ?, ?)",
            (int(activity_id), int(member_id), clean_text(role)),
        )


def delete_activity(activity_id: int) -> None:
    execute("DELETE FROM activity_members WHERE activity_id=?", (int(activity_id),))
    execute("DELETE FROM activities WHERE id=?", (int(activity_id),))


def activities_df(
    month: str = "",
    project_id: Optional[int] = None,
    partner_id: Optional[int] = None,
    person_id: Optional[int] = None,
    status: str = "Tous",
) -> pd.DataFrame:
    where = []
    params: List = []
    if month:
        where.append("(a.period_month = ? OR substr(a.planned_date,1,7)=? OR substr(a.deadline,1,7)=?)")
        params.extend([month, month, month])
    if project_id:
        where.append("a.project_id = ?")
        params.append(project_id)
    if partner_id:
        where.append("a.partner_id = ?")
        params.append(partner_id)
    if person_id:
        where.append("(a.responsible_id = ? OR a.id IN (SELECT activity_id FROM activity_members WHERE member_id = ?))")
        params.extend([person_id, person_id])
    if status != "Tous":
        where.append("a.status = ?")
        params.append(status)
    where_sql = "WHERE " + " AND ".join(where) if where else ""
    sql = f"""
        SELECT
            a.id AS ID,
            a.title AS Activité,
            a.description AS Description,
            COALESCE(p.name, '') AS Projet,
            COALESCE(pa.name, '') AS Partenaire,
            a.period_month AS "Mois prévu",
            a.planned_date AS "Jour prévu",
            a.deadline AS "Date limite",
            a.priority AS Priorité,
            a.status AS Statut,
            COALESCE(resp.name, '') AS "Responsable principal",
            COALESCE(resp.poste, '') AS "Poste responsable",
            COALESCE(a.location, '') AS Lieu,
            COALESCE(a.expected_output, '') AS "Résultat attendu",
            COALESCE(a.notes, '') AS Notes,
            COALESCE(GROUP_CONCAT(
                CASE
                    WHEN mem.name IS NULL THEN NULL
                    WHEN am.role_in_activity IS NOT NULL AND am.role_in_activity <> '' THEN mem.name || ' (' || am.role_in_activity || ')'
                    ELSE mem.name
                END, '; '
            ), '') AS "Membres affectés"
        FROM activities a
        LEFT JOIN projects p ON p.id = a.project_id
        LEFT JOIN partners pa ON pa.id = a.partner_id
        LEFT JOIN team_members resp ON resp.id = a.responsible_id
        LEFT JOIN activity_members am ON am.activity_id = a.id
        LEFT JOIN team_members mem ON mem.id = am.member_id
        {where_sql}
        GROUP BY a.id
        ORDER BY COALESCE(a.planned_date, a.deadline, a.period_month), a.priority DESC, a.id DESC
    """
    df = qdf(sql, tuple(params))
    if df.empty:
        return df
    df["Urgence auto"] = df.apply(urgency_for_row, axis=1)
    df["Jours avant limite"] = df.apply(days_to_deadline, axis=1)
    return df


def get_activity(activity_id: int) -> Optional[sqlite3.Row]:
    return fetch_one("SELECT * FROM activities WHERE id=?", (int(activity_id),))


def get_activity_members(activity_id: int) -> pd.DataFrame:
    return qdf(
        """
        SELECT tm.id, tm.name, tm.poste, COALESCE(am.role_in_activity, '') AS role_in_activity
        FROM activity_members am
        JOIN team_members tm ON tm.id = am.member_id
        WHERE am.activity_id = ?
        ORDER BY tm.name
        """,
        (int(activity_id),),
    )


# -----------------------------------------------------------------------------
# Import / export
# -----------------------------------------------------------------------------

def read_uploaded_file(uploaded_file) -> Dict[str, pd.DataFrame]:
    suffix = Path(uploaded_file.name).suffix.lower()
    if suffix in [".xlsx", ".xls"]:
        return pd.read_excel(uploaded_file, sheet_name=None)
    if suffix == ".csv":
        return {"CSV": pd.read_csv(uploaded_file)}
    if suffix == ".docx":
        if Document is None:
            raise RuntimeError("python-docx n'est pas installé. Ajoutez python-docx dans requirements.txt")
        doc = Document(uploaded_file)
        tables = {}
        for i, table in enumerate(doc.tables, start=1):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            header = rows[0]
            data = rows[1:] if len(rows) > 1 else []
            tables[f"Tableau {i}"] = pd.DataFrame(data, columns=header)
        if tables:
            return tables
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return {"Paragraphes": pd.DataFrame({"Activité": paragraphs})}
    raise ValueError("Format non supporté. Utilisez Excel (.xlsx), CSV ou Word (.docx).")


def import_team_df(df: pd.DataFrame) -> int:
    mapping = build_column_map(df.columns)
    count = 0
    for _, row in df.iterrows():
        name = clean_text(get_col_value(row, {"name": mapping.get("title") or mapping.get("responsible") or "Nom"}, "name"))
        if not name:
            possible = [c for c in df.columns if normalize(c) in ["nom", "name", "membre"]]
            if possible:
                name = clean_text(row.get(possible[0]))
        if not name:
            continue
        poste_col = next((c for c in df.columns if normalize(c) in ["poste", "fonction", "role", "position"]), None)
        email_col = next((c for c in df.columns if normalize(c) in ["email", "mail", "courriel"]), None)
        phone_col = next((c for c in df.columns if normalize(c) in ["telephone", "phone", "tel"]), None)
        active_col = next((c for c in df.columns if normalize(c) in ["actif", "active"]), None)
        notes_col = next((c for c in df.columns if normalize(c) in ["notes", "observation", "observations"]), None)
        execute(
            """
            INSERT OR IGNORE INTO team_members(name, poste, email, phone, active, notes)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (
                name,
                clean_text(row.get(poste_col, "")) if poste_col else "",
                clean_text(row.get(email_col, "")) if email_col else "",
                clean_text(row.get(phone_col, "")) if phone_col else "",
                0 if active_col and normalize(row.get(active_col, "oui")) in ["non", "no", "0", "false"] else 1,
                clean_text(row.get(notes_col, "")) if notes_col else "",
            ),
        )
        count += 1
    return count


def import_simple_list(df: pd.DataFrame, kind: str) -> int:
    name_candidates = ["projet", "project"] if kind == "project" else ["partenaire", "partner"]
    name_col = next((c for c in df.columns if normalize(c) in name_candidates), None)
    if not name_col and len(df.columns):
        name_col = df.columns[0]
    desc_col = next((c for c in df.columns if normalize(c) in ["description", "details", "detail"]), None)
    count = 0
    for _, row in df.iterrows():
        name = clean_text(row.get(name_col, "")) if name_col else ""
        desc = clean_text(row.get(desc_col, "")) if desc_col else ""
        if not name:
            continue
        if kind == "project":
            ensure_project(name, desc)
        else:
            ensure_partner(name, desc)
        count += 1
    return count


def import_activities_df(
    df: pd.DataFrame,
    default_month: str,
    default_project_id: Optional[int],
    default_partner_id: Optional[int],
    create_missing_members: bool,
    source_file: str,
) -> Tuple[int, List[str]]:
    mapping = build_column_map(df.columns)
    errors = []
    count = 0
    default_project = fetch_one("SELECT name FROM projects WHERE id=?", (default_project_id,)) if default_project_id else None
    default_partner = fetch_one("SELECT name FROM partners WHERE id=?", (default_partner_id,)) if default_partner_id else None
    default_project_name = default_project["name"] if default_project else ""
    default_partner_name = default_partner["name"] if default_partner else ""
    for idx, row in df.iterrows():
        title = clean_text(get_col_value(row, mapping, "title", ""))
        if not title:
            non_empty = [clean_text(v) for v in row.tolist() if clean_text(v)]
            title = non_empty[0] if non_empty else ""
        if not title:
            continue
        project_name = clean_text(get_col_value(row, mapping, "project", default_project_name)) or default_project_name
        partner_name = clean_text(get_col_value(row, mapping, "partner", default_partner_name)) or default_partner_name
        project_id = ensure_project(project_name) if project_name else default_project_id
        partner_id = ensure_partner(partner_name) if partner_name else default_partner_id
        planned_date = parse_date_value(get_col_value(row, mapping, "planned_date", ""))
        deadline = parse_date_value(get_col_value(row, mapping, "deadline", ""))
        period_month = parse_month_value(get_col_value(row, mapping, "period_month", ""), planned_date or deadline, default_month)
        priority = normalize_priority(get_col_value(row, mapping, "priority", "Moyenne"))
        status = normalize_status(get_col_value(row, mapping, "status", "À faire"))
        responsible_name = clean_text(get_col_value(row, mapping, "responsible", ""))
        responsible_id = ensure_member(responsible_name) if create_missing_members and responsible_name else None
        member_roles: Dict[int, str] = {}
        for member_name in split_members(get_col_value(row, mapping, "members", "")):
            if create_missing_members:
                member_id = ensure_member(member_name)
            else:
                member = fetch_one("SELECT id FROM team_members WHERE name=?", (member_name,))
                member_id = int(member["id"]) if member else None
            if member_id:
                member_roles[int(member_id)] = ""
        try:
            add_activity(
                title=title,
                description=clean_text(get_col_value(row, mapping, "description", "")),
                project_id=project_id,
                partner_id=partner_id,
                period_month=period_month,
                planned_date=planned_date,
                deadline=deadline,
                priority=priority,
                status=status,
                responsible_id=responsible_id,
                location=clean_text(get_col_value(row, mapping, "location", "")),
                expected_output=clean_text(get_col_value(row, mapping, "expected_output", "")),
                notes=clean_text(get_col_value(row, mapping, "notes", "")),
                member_roles=member_roles,
                source_file=source_file,
            )
            count += 1
        except Exception as exc:
            errors.append(f"Ligne {idx + 2}: {exc}")
    return count, errors


def export_excel_bytes(filtered_df: pd.DataFrame) -> bytes:
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        filtered_df.to_excel(writer, index=False, sheet_name="Planning filtre")
        list_team().to_excel(writer, index=False, sheet_name="Equipe")
        list_projects().to_excel(writer, index=False, sheet_name="Projets")
        list_partners().to_excel(writer, index=False, sheet_name="Partenaires")
        all_df = activities_df()
        all_df.to_excel(writer, index=False, sheet_name="Toutes activites")
    output.seek(0)
    return output.read()


def db_backup_bytes() -> bytes:
    # Force a checkpoint-like copy by reading the database file.
    if DB_PATH.exists():
        return DB_PATH.read_bytes()
    return b""


# -----------------------------------------------------------------------------
# Filters
# -----------------------------------------------------------------------------

def sidebar_filters() -> Tuple[str, Optional[int], Optional[int], Optional[int], str]:
    st.sidebar.markdown("### Filtres du dashboard")
    selected_day = st.sidebar.date_input("Mois à analyser", value=date.today().replace(day=1))
    selected_month = selected_day.strftime("%Y-%m")

    projects = list_projects(active_only=True)
    project_options = {"Tous les projets": None}
    project_options.update({row["name"]: int(row["id"]) for _, row in projects.iterrows()})
    project_label = st.sidebar.selectbox("Projet", list(project_options.keys()))

    partners = list_partners()
    partner_options = {"Tous les partenaires": None}
    partner_options.update({row["name"]: int(row["id"]) for _, row in partners.iterrows()})
    partner_label = st.sidebar.selectbox("Partenaire", list(partner_options.keys()))

    team = list_team(active_only=True)
    person_options = {"Toutes les personnes": None}
    person_options.update({f"{row['name']} — {row['poste']}" if row["poste"] else row["name"]: int(row["id"]) for _, row in team.iterrows()})
    person_label = st.sidebar.selectbox("Personne", list(person_options.keys()))

    status_label = st.sidebar.selectbox("Statut", ["Tous"] + STATUSES)
    return selected_month, project_options[project_label], partner_options[partner_label], person_options[person_label], status_label


# -----------------------------------------------------------------------------
# Pages
# -----------------------------------------------------------------------------

def page_dashboard() -> None:
    selected_month, project_id, partner_id, person_id, status = sidebar_filters()
    df = activities_df(selected_month, project_id, partner_id, person_id, status)

    st.subheader(f"Dashboard — {month_label(selected_month)}")
    if df.empty:
        st.info("Aucune activité trouvée pour ces filtres. Ajoutez des activités ou importez un fichier Excel/Word.")
        return

    urgent_count = int(df["Urgence auto"].isin(["Urgent", "En retard"]).sum())
    late_count = int((df["Urgence auto"] == "En retard").sum())
    in_progress_count = int((df["Statut"] == "En cours").sum())
    done_count = int((df["Statut"] == "Terminé").sum())

    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        kpi_card("Activités du mois", str(len(df)), "Filtrées selon vos critères")
    with col2:
        kpi_card("Urgentes / en retard", str(urgent_count), "Priorité haute, critique ou échéance proche")
    with col3:
        kpi_card("En retard", str(late_count), "Date limite dépassée")
    with col4:
        kpi_card("En cours", str(in_progress_count), "Activités démarrées")
    with col5:
        kpi_card("Terminées", str(done_count), "Activités clôturées")

    st.markdown("### Planning mensuel")
    display_cols = [
        "Urgence auto", "Projet", "Partenaire", "Activité", "Jour prévu", "Date limite", "Priorité",
        "Statut", "Responsable principal", "Poste responsable", "Membres affectés", "Lieu", "Résultat attendu"
    ]
    st.dataframe(df[display_cols], use_container_width=True, hide_index=True)

    expander = st.expander("Voir les notes et descriptions")
    with expander:
        st.dataframe(df[["ID", "Activité", "Description", "Notes"]], use_container_width=True, hide_index=True)

    st.markdown("### Analyse rapide")
    chart_col1, chart_col2 = st.columns(2)
    with chart_col1:
        by_project = df.groupby("Projet", dropna=False).size().reset_index(name="Nombre")
        chart = alt.Chart(by_project).mark_bar().encode(
            x=alt.X("Nombre:Q", title="Activités"),
            y=alt.Y("Projet:N", sort="-x", title="Projet"),
            tooltip=["Projet", "Nombre"],
        ).properties(height=280)
        st.altair_chart(chart, use_container_width=True)
    with chart_col2:
        by_responsible = df.groupby("Responsable principal", dropna=False).size().reset_index(name="Nombre")
        by_responsible["Responsable principal"] = by_responsible["Responsable principal"].replace("", "Non assigné")
        chart = alt.Chart(by_responsible).mark_bar().encode(
            x=alt.X("Nombre:Q", title="Activités"),
            y=alt.Y("Responsable principal:N", sort="-x", title="Responsable"),
            tooltip=["Responsable principal", "Nombre"],
        ).properties(height=280)
        st.altair_chart(chart, use_container_width=True)

    st.markdown("### Vue par jour")
    calendar_df = df.copy()
    calendar_df["Jour"] = calendar_df["Jour prévu"].replace("", pd.NA).fillna(calendar_df["Date limite"])
    calendar_df = calendar_df[calendar_df["Jour"].notna() & (calendar_df["Jour"] != "")]
    if calendar_df.empty:
        st.caption("Aucune date précise. Utilisez la page Activités pour ajouter un jour prévu.")
    else:
        calendar_df = calendar_df.sort_values("Jour")
        st.dataframe(
            calendar_df[["Jour", "Urgence auto", "Activité", "Projet", "Responsable principal", "Membres affectés", "Statut"]],
            use_container_width=True,
            hide_index=True,
        )

    st.download_button(
        "⬇️ Exporter ce planning filtré en Excel",
        data=export_excel_bytes(df),
        file_name=f"SADABE_PLANIFIUS_{selected_month}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


def select_project(label: str, default_id: Optional[int] = None, key: Optional[str] = None) -> Optional[int]:
    projects = list_projects(active_only=True)
    if projects.empty:
        st.warning("Ajoutez au moins un projet.")
        return None
    options = {row["name"]: int(row["id"]) for _, row in projects.iterrows()}
    labels = list(options.keys())
    index = labels.index(next((k for k, v in options.items() if v == default_id), labels[0])) if labels else 0
    return options[st.selectbox(label, labels, index=index, key=key)]


def select_partner(label: str, default_id: Optional[int] = None, allow_none: bool = True, key: Optional[str] = None) -> Optional[int]:
    partners = list_partners()
    options = {"Aucun": None} if allow_none else {}
    options.update({row["name"]: int(row["id"]) for _, row in partners.iterrows()})
    labels = list(options.keys())
    default_label = next((k for k, v in options.items() if v == default_id), labels[0])
    return options[st.selectbox(label, labels, index=labels.index(default_label), key=key)]


def select_member(label: str, default_id: Optional[int] = None, allow_none: bool = True, key: Optional[str] = None) -> Optional[int]:
    team = list_team(active_only=True)
    options = {"Aucun": None} if allow_none else {}
    options.update({f"{row['name']} — {row['poste']}" if row["poste"] else row["name"]: int(row["id"]) for _, row in team.iterrows()})
    labels = list(options.keys())
    default_label = next((k for k, v in options.items() if v == default_id), labels[0])
    return options[st.selectbox(label, labels, index=labels.index(default_label), key=key)]


def multiselect_members(default_ids: Optional[List[int]] = None, key: Optional[str] = None) -> List[int]:
    default_ids = default_ids or []
    team = list_team(active_only=True)
    options = {f"{row['name']} — {row['poste']}" if row["poste"] else row["name"]: int(row["id"]) for _, row in team.iterrows()}
    default_labels = [label for label, mid in options.items() if mid in default_ids]
    selected_labels = st.multiselect("Membres affectés à l'activité", list(options.keys()), default=default_labels, key=key)
    return [options[label] for label in selected_labels]


def page_activities() -> None:
    st.subheader("Activités et planification détaillée")
    st.caption("Ajoutez une activité, assignez un responsable principal, puis ajoutez les membres SADABE qui participeront.")

    tab_add, tab_edit, tab_table = st.tabs(["➕ Ajouter", "✏️ Modifier", "📋 Liste"])

    with tab_add:
        with st.form("add_activity_form", clear_on_submit=True):
            title = st.text_input("Activité / tâche *")
            description = st.text_area("Description")
            col1, col2 = st.columns(2)
            with col1:
                project_id = select_project("Projet *", key="add_project")
                partner_id = select_partner("Partenaire", key="add_partner")
                selected_month = st.date_input("Mois prévu", value=date.today().replace(day=1)).strftime("%Y-%m")
            with col2:
                responsible_id = select_member("Responsable principal", key="add_responsible")
                priority = st.selectbox("Priorité", PRIORITIES, index=1)
                status = st.selectbox("Statut", STATUSES, index=0)

            col3, col4 = st.columns(2)
            with col3:
                use_planned = st.checkbox("Ajouter un jour précis", value=True)
                planned_date = st.date_input("Jour prévu", value=date.today()).isoformat() if use_planned else ""
            with col4:
                use_deadline = st.checkbox("Ajouter une date limite", value=True)
                deadline = st.date_input("Date limite", value=date.today()).isoformat() if use_deadline else ""

            location = st.text_input("Lieu / site")
            expected_output = st.text_input("Résultat attendu / livrable")
            selected_member_ids = multiselect_members(key="add_members")
            member_roles = {}
            if selected_member_ids:
                st.markdown("**Rôle des membres dans cette activité**")
                team = list_team(active_only=True).set_index("id")
                for mid in selected_member_ids:
                    name = team.loc[mid, "name"] if mid in team.index else str(mid)
                    member_roles[mid] = st.text_input(f"Rôle de {name}", value="", key=f"role_add_{mid}")
            notes = st.text_area("Notes")
            submitted = st.form_submit_button("Enregistrer l'activité")
            if submitted:
                if not title.strip():
                    st.error("Le titre de l'activité est obligatoire.")
                else:
                    add_activity(
                        title=title,
                        description=description,
                        project_id=project_id,
                        partner_id=partner_id,
                        period_month=selected_month,
                        planned_date=planned_date,
                        deadline=deadline,
                        priority=priority,
                        status=status,
                        responsible_id=responsible_id,
                        location=location,
                        expected_output=expected_output,
                        notes=notes,
                        member_roles=member_roles,
                    )
                    st.success("Activité ajoutée.")
                    st.rerun()

    with tab_edit:
        all_activities = activities_df()
        if all_activities.empty:
            st.info("Aucune activité à modifier.")
        else:
            option_map = {f"#{row['ID']} — {row['Activité']}": int(row["ID"]) for _, row in all_activities.iterrows()}
            label = st.selectbox("Choisir une activité", list(option_map.keys()))
            activity_id = option_map[label]
            activity = get_activity(activity_id)
            activity_members = get_activity_members(activity_id)
            default_member_ids = activity_members["id"].astype(int).tolist() if not activity_members.empty else []

            with st.form("edit_activity_form"):
                title = st.text_input("Activité / tâche *", value=activity["title"])
                description = st.text_area("Description", value=activity["description"] or "")
                col1, col2 = st.columns(2)
                with col1:
                    project_id = select_project("Projet *", activity["project_id"], key="edit_project")
                    partner_id = select_partner("Partenaire", activity["partner_id"], key="edit_partner")
                    month_value = activity["period_month"] or date.today().strftime("%Y-%m")
                    try:
                        month_date = datetime.strptime(month_value + "-01", "%Y-%m-%d").date()
                    except Exception:
                        month_date = date.today().replace(day=1)
                    period_month = st.date_input("Mois prévu", value=month_date, key="edit_month").strftime("%Y-%m")
                with col2:
                    responsible_id = select_member("Responsable principal", activity["responsible_id"], key="edit_responsible")
                    priority = st.selectbox("Priorité", PRIORITIES, index=PRIORITIES.index(activity["priority"]) if activity["priority"] in PRIORITIES else 1)
                    status = st.selectbox("Statut", STATUSES, index=STATUSES.index(activity["status"]) if activity["status"] in STATUSES else 0)
                col3, col4 = st.columns(2)
                with col3:
                    has_planned = bool(activity["planned_date"])
                    use_planned = st.checkbox("Ajouter un jour précis", value=has_planned, key="edit_use_planned")
                    planned_default = datetime.fromisoformat(activity["planned_date"]).date() if has_planned else date.today()
                    planned_date = st.date_input("Jour prévu", value=planned_default, key="edit_planned").isoformat() if use_planned else ""
                with col4:
                    has_deadline = bool(activity["deadline"])
                    use_deadline = st.checkbox("Ajouter une date limite", value=has_deadline, key="edit_use_deadline")
                    deadline_default = datetime.fromisoformat(activity["deadline"]).date() if has_deadline else date.today()
                    deadline = st.date_input("Date limite", value=deadline_default, key="edit_deadline").isoformat() if use_deadline else ""
                location = st.text_input("Lieu / site", value=activity["location"] or "")
                expected_output = st.text_input("Résultat attendu / livrable", value=activity["expected_output"] or "")
                selected_member_ids = multiselect_members(default_member_ids, key="edit_members")
                member_roles = {}
                current_roles = {int(r["id"]): r["role_in_activity"] for _, r in activity_members.iterrows()} if not activity_members.empty else {}
                if selected_member_ids:
                    st.markdown("**Rôle des membres dans cette activité**")
                    team = list_team(active_only=True).set_index("id")
                    for mid in selected_member_ids:
                        name = team.loc[mid, "name"] if mid in team.index else str(mid)
                        member_roles[mid] = st.text_input(f"Rôle de {name}", value=current_roles.get(mid, ""), key=f"role_edit_{mid}")
                notes = st.text_area("Notes", value=activity["notes"] or "")
                col_save, col_delete = st.columns(2)
                save = col_save.form_submit_button("Sauvegarder les modifications")
                delete = col_delete.form_submit_button("Supprimer cette activité")
                if save:
                    update_activity(
                        activity_id,
                        {
                            "title": title,
                            "description": description,
                            "project_id": project_id,
                            "partner_id": partner_id,
                            "period_month": period_month,
                            "planned_date": planned_date,
                            "deadline": deadline,
                            "priority": priority,
                            "status": status,
                            "responsible_id": responsible_id,
                            "location": location,
                            "expected_output": expected_output,
                            "notes": notes,
                        },
                        member_roles,
                    )
                    st.success("Activité modifiée.")
                    st.rerun()
                if delete:
                    delete_activity(activity_id)
                    st.warning("Activité supprimée.")
                    st.rerun()

    with tab_table:
        selected_month, project_id, partner_id, person_id, status = sidebar_filters()
        df = activities_df(selected_month, project_id, partner_id, person_id, status)
        st.dataframe(df, use_container_width=True, hide_index=True)


def page_team() -> None:
    st.subheader("Équipe SADABE")
    st.caption("Ajoutez les membres de l'équipe et leur poste. Ces personnes peuvent ensuite être responsables ou membres d'une activité.")

    with st.form("team_add_form", clear_on_submit=True):
        col1, col2 = st.columns(2)
        with col1:
            name = st.text_input("Nom du membre *")
            poste = st.text_input("Poste / fonction")
        with col2:
            email = st.text_input("Email")
            phone = st.text_input("Téléphone")
        notes = st.text_area("Notes")
        submitted = st.form_submit_button("Ajouter le membre")
        if submitted:
            if not name.strip():
                st.error("Le nom est obligatoire.")
            else:
                try:
                    execute(
                        "INSERT INTO team_members(name, poste, email, phone, notes, active) VALUES (?, ?, ?, ?, ?, 1)",
                        (name.strip(), poste.strip(), email.strip(), phone.strip(), notes.strip()),
                    )
                    st.success("Membre ajouté.")
                    st.rerun()
                except sqlite3.IntegrityError:
                    st.error("Ce membre existe déjà.")

    st.markdown("### Liste de l'équipe")
    team = list_team()
    if team.empty:
        st.info("Aucun membre enregistré.")
    else:
        editable = team.rename(columns={"name": "Nom", "poste": "Poste", "email": "Email", "phone": "Téléphone", "active": "Actif", "notes": "Notes"})
        editable["Actif"] = editable["Actif"].astype(bool)
        edited = st.data_editor(editable, use_container_width=True, hide_index=True, disabled=["id"])
        if st.button("Sauvegarder les modifications de l'équipe"):
            try:
                for _, row in edited.iterrows():
                    execute(
                        "UPDATE team_members SET name=?, poste=?, email=?, phone=?, active=?, notes=? WHERE id=?",
                        (
                            clean_text(row["Nom"]),
                            clean_text(row["Poste"]),
                            clean_text(row["Email"]),
                            clean_text(row["Téléphone"]),
                            1 if bool(row["Actif"]) else 0,
                            clean_text(row["Notes"]),
                            int(row["id"]),
                        ),
                    )
                st.success("Équipe mise à jour.")
                st.rerun()
            except sqlite3.IntegrityError:
                st.error("Deux membres ne peuvent pas avoir exactement le même nom.")


def page_projects_partners() -> None:
    st.subheader("Projets et partenaires")
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("### Ajouter un projet")
        with st.form("project_form", clear_on_submit=True):
            name = st.text_input("Nom du projet")
            desc = st.text_area("Description")
            if st.form_submit_button("Ajouter le projet"):
                if name.strip():
                    ensure_project(name, desc)
                    st.success("Projet ajouté.")
                    st.rerun()
                else:
                    st.error("Nom obligatoire.")
    with col2:
        st.markdown("### Ajouter un partenaire")
        with st.form("partner_form", clear_on_submit=True):
            name = st.text_input("Nom du partenaire")
            desc = st.text_area("Description")
            if st.form_submit_button("Ajouter le partenaire"):
                if name.strip():
                    ensure_partner(name, desc)
                    st.success("Partenaire ajouté.")
                    st.rerun()
                else:
                    st.error("Nom obligatoire.")

    tab_projects, tab_partners = st.tabs(["Projets", "Partenaires"])
    with tab_projects:
        projects = list_projects().rename(columns={"name": "Projet", "description": "Description", "active": "Actif"})
        projects["Actif"] = projects["Actif"].astype(bool)
        edited_projects = st.data_editor(projects, use_container_width=True, hide_index=True, disabled=["id"])
        if st.button("Sauvegarder les projets"):
            try:
                for _, row in edited_projects.iterrows():
                    execute(
                        "UPDATE projects SET name=?, description=?, active=? WHERE id=?",
                        (clean_text(row["Projet"]), clean_text(row["Description"]), 1 if bool(row["Actif"]) else 0, int(row["id"])),
                    )
                st.success("Projets mis à jour.")
                st.rerun()
            except sqlite3.IntegrityError:
                st.error("Deux projets ne peuvent pas avoir exactement le même nom.")
    with tab_partners:
        partners = list_partners().rename(columns={"name": "Partenaire", "description": "Description"})
        edited_partners = st.data_editor(partners, use_container_width=True, hide_index=True, disabled=["id"])
        if st.button("Sauvegarder les partenaires"):
            try:
                for _, row in edited_partners.iterrows():
                    execute(
                        "UPDATE partners SET name=?, description=? WHERE id=?",
                        (clean_text(row["Partenaire"]), clean_text(row["Description"]), int(row["id"])),
                    )
                st.success("Partenaires mis à jour.")
                st.rerun()
            except sqlite3.IntegrityError:
                st.error("Deux partenaires ne peuvent pas avoir exactement le même nom.")


def page_import_export() -> None:
    st.subheader("Import / Export")
    st.caption("Importez des activités depuis Excel, CSV ou Word. Le modèle Excel fourni contient les colonnes recommandées.")

    template_path = SAMPLE_DIR / "modele_import_sadabe_planifius.xlsx"
    if template_path.exists():
        st.download_button(
            "⬇️ Télécharger le modèle Excel d'import",
            data=template_path.read_bytes(),
            file_name="modele_import_sadabe_planifius.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    uploaded = st.file_uploader("Importer un fichier Excel, CSV ou Word", type=["xlsx", "xls", "csv", "docx"])
    if uploaded is not None:
        try:
            sheets = read_uploaded_file(uploaded)
        except Exception as exc:
            st.error(str(exc))
            return
        sheet_name = st.selectbox("Feuille / tableau à importer comme activités", list(sheets.keys()))
        df = sheets[sheet_name]
        st.markdown("### Aperçu")
        st.dataframe(df.head(20), use_container_width=True)

        col1, col2, col3 = st.columns(3)
        with col1:
            default_month = st.date_input("Mois par défaut", value=date.today().replace(day=1), key="import_month").strftime("%Y-%m")
        with col2:
            default_project_id = select_project("Projet par défaut", key="import_default_project")
        with col3:
            default_partner_id = select_partner("Partenaire par défaut", key="import_default_partner")
        create_missing_members = st.checkbox("Créer automatiquement les responsables/membres inconnus", value=True)
        import_other_tabs = st.checkbox("Si le fichier contient les onglets Equipe, Projets ou Partenaires, les importer aussi", value=True)

        if st.button("Importer dans SADABE PLANIFIUS"):
            total_other = []
            if import_other_tabs:
                for name, tab_df in sheets.items():
                    norm = normalize(name)
                    if norm in ["equipe", "team", "membres"]:
                        total_other.append(f"Équipe: {import_team_df(tab_df)}")
                    elif norm in ["projets", "projects"]:
                        total_other.append(f"Projets: {import_simple_list(tab_df, 'project')}")
                    elif norm in ["partenaires", "partners"]:
                        total_other.append(f"Partenaires: {import_simple_list(tab_df, 'partner')}")
            count, errors = import_activities_df(
                df,
                default_month=default_month,
                default_project_id=default_project_id,
                default_partner_id=default_partner_id,
                create_missing_members=create_missing_members,
                source_file=uploaded.name,
            )
            st.success(f"{count} activité(s) importée(s)." + (" Autres onglets: " + ", ".join(total_other) if total_other else ""))
            if errors:
                st.warning("Certaines lignes n'ont pas été importées.")
                st.write(errors[:20])
            st.rerun()

    st.markdown("---")
    st.markdown("### Export")
    all_df = activities_df()
    st.download_button(
        "⬇️ Exporter toute la base en Excel",
        data=export_excel_bytes(all_df),
        file_name="SADABE_PLANIFIUS_export_complet.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    if DB_PATH.exists():
        st.download_button(
            "⬇️ Sauvegarder la base SQLite",
            data=db_backup_bytes(),
            file_name="sadabe_planifius.db",
            mime="application/octet-stream",
        )


def page_about() -> None:
    st.subheader("À propos")
    st.markdown(
        """
        **SADABE PLANIFIUS** est une application open source de planification mensuelle pour SADABE.

        Elle permet de :
        - voir les activités prévues par mois ;
        - filtrer le dashboard par projet, partenaire, personne et statut ;
        - ajouter les responsables principaux et les membres affectés ;
        - gérer les postes des membres de l'équipe ;
        - importer des fichiers Excel, CSV ou Word ;
        - exporter la planification filtrée en Excel.

        **Projets préinstallés :** SOS Lemurs, Darwin Initiatives, Seacology, Rainforest Trust.  
        **Partenaires préinstallés :** TGBS (MBG), MfM, UWE, Regen, UNI, ENS.
        """
    )
    if LOGO_PATH.exists():
        st.image(str(LOGO_PATH), width=180)


# -----------------------------------------------------------------------------
# Main
# -----------------------------------------------------------------------------

def main() -> None:
    st.set_page_config(page_title=APP_NAME, page_icon="🌿", layout="wide")
    inject_css()
    get_connection()
    render_header()

    st.sidebar.image(str(LOGO_PATH), width=120) if LOGO_PATH.exists() else None
    page = st.sidebar.radio(
        "Navigation",
        ["Dashboard", "Activités", "Équipe SADABE", "Projets & partenaires", "Import / Export", "À propos"],
    )

    if page == "Dashboard":
        page_dashboard()
    elif page == "Activités":
        page_activities()
    elif page == "Équipe SADABE":
        page_team()
    elif page == "Projets & partenaires":
        page_projects_partners()
    elif page == "Import / Export":
        page_import_export()
    else:
        page_about()


if __name__ == "__main__":
    main()
